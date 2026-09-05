"""Render the anonymous competition submission from reproducible public-source outputs.

The document deliberately distinguishes public accounting observations from tax-risk
conclusions.  It does not manufacture tax filings, contracts, interviews, or fieldwork.
"""

from __future__ import annotations

import csv
import os
import re
import subprocess
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from taxrisk.workflow import install_script_logging

ROOT = Path(__file__).resolve().parents[1]
FIELDS = ROOT / "outputs" / "tables" / "yonyou_public_financial_fields.csv"
INDICATORS = ROOT / "outputs" / "tables" / "yonyou_public_indicator_observations.csv"
POLICIES = ROOT / "sources" / "policies" / "index.csv"
REPORT_TEXTS = {
    2023: ROOT / "sources" / "curated" / "open_data" / "yonyou" / "2023_annual_report.txt",
    2024: ROOT / "sources" / "curated" / "open_data" / "yonyou" / "2024_annual_report.txt",
}
OUT = ROOT / "outputs" / "submission"
DOCX = OUT / "甲软件和信息技术服务企业税收风险管控案例.docx"
PDF = OUT / "甲软件和信息技术服务企业税收风险管控案例.pdf"
FONT_CONFIG = ROOT / "config" / "fontconfig.conf"

CN = "宋体"
KAI = "楷体"
LIGHT_BLUE = "D9EAF7"


def load_rows(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def format_million(value: float) -> str:
    return f"{value / 1_000_000:,.2f}"


def format_pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def extract_note_amount(text: str, section: str, label: str) -> int:
    start = text.find(section)
    if start < 0:
        raise ValueError(f"missing section: {section}")
    excerpt = text[start : start + 7000]
    match = re.search(rf"{re.escape(label)}\s+([\d,]+)", excerpt)
    if not match:
        raise ValueError(f"missing {label} in {section}")
    return int(match.group(1).replace(",", ""))


def collect_data() -> dict[str, object]:
    fields = load_rows(FIELDS)
    indicators = load_rows(INDICATORS)
    policies = {row["policy_id"]: row for row in load_rows(POLICIES)}
    amounts: dict[int, dict[str, float]] = defaultdict(dict)
    locations: dict[int, dict[str, str]] = defaultdict(dict)
    for row in fields:
        period = int(row["period"])
        amounts[period][row["field"]] = float(row["amount_cny"])
        locations[period][row["field"]] = (
            f"[S{period - 2022}]第{row['report_page']}页，{row['source_section']}"
        )
    by_indicator: dict[str, dict[int, dict[str, str]]] = defaultdict(dict)
    for row in indicators:
        by_indicator[row["indicator_id"]][int(row["period"])] = row
    tax_notes: dict[int, dict[str, int]] = {}
    for year, source in REPORT_TEXTS.items():
        text = source.read_text(encoding="utf-8")
        tax_notes[year] = {
            "vat_payable": extract_note_amount(text, "40、 应交税费", "增值税"),
            "cit_payable": extract_note_amount(text, "40、 应交税费", "企业所得税"),
            "iit_payable": extract_note_amount(text, "40、 应交税费", "个人所得税"),
            "stamp_tax_expense": extract_note_amount(text, "62、 税金及附加", "印花税"),
        }
    return {
        "amounts": amounts,
        "locations": locations,
        "indicators": by_indicator,
        "tax_notes": tax_notes,
        "policies": policies,
    }


def set_run_font(run, name: str = CN, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Times New Roman"
    rfonts = run._element.rPr.rFonts
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def set_cell_text(cell, text: str, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        set_cell_shading(cell, fill)


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, fill=LIGHT_BLUE)
        if widths:
                table.rows[0].cells[index].width = Cm(widths[index])
    repeat_table_header(table.rows[0])
    keep_table_row_together(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            if widths:
                cells[index].width = Cm(widths[index])
        keep_table_row_together(table.rows[-1])
    doc.add_paragraph()


def add_para(doc: Document, text: str = "", *, align=None, bold: bool = False, size: float = 10.5, font: str = CN, color: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.3
    paragraph.paragraph_format.space_before = Pt(3.4)
    paragraph.paragraph_format.space_after = Pt(3.4)
    run = paragraph.add_run(text)
    set_run_font(run, name=font, size=size, bold=bold, color=color)


def add_heading(doc: Document, level: int, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.3
    paragraph.paragraph_format.space_before = Pt(3.4)
    paragraph.paragraph_format.space_after = Pt(3.4)
    run = paragraph.add_run(text)
    set_run_font(run, name=CN, size=12, bold=True)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = 1.3
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    set_run_font(run, size=8.5)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)
    run2 = footer.add_run(" 页")
    set_run_font(run2, size=8.5)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)
    add_page_number(section)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal_fonts = normal._element.rPr.rFonts
    normal_fonts.set(qn("w:ascii"), "Times New Roman")
    normal_fonts.set(qn("w:hAnsi"), "Times New Roman")
    normal_fonts.set(qn("w:eastAsia"), CN)
    normal.font.size = Pt(10.5)
    for style_name in ("Title", "Subtitle"):
        style = doc.styles[style_name]
        style.font.name = KAI
        style._element.rPr.rFonts.set(qn("w:eastAsia"), KAI)


def add_cover(doc: Document) -> None:
    for _ in range(7):
        add_para(doc, "")
    add_para(doc, "甲软件和信息技术服务企业", align=WD_ALIGN_PARAGRAPH.CENTER, font=KAI, size=15, bold=True)
    add_para(doc, "税收风险管控案例", align=WD_ALIGN_PARAGRAPH.CENTER, font=KAI, size=15, bold=True)
    add_para(doc, "基于两期公开年报与官方政策的可复现案头研究", align=WD_ALIGN_PARAGRAPH.CENTER, font=CN, size=12)
    for _ in range(8):
        add_para(doc, "")
    add_para(doc, "提交文本（匿名版）", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_para(doc, "2026年9月", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    doc.add_page_break()


def build_document(data: dict[str, object]) -> Document:
    amounts = data["amounts"]
    locations = data["locations"]
    indicators = data["indicators"]
    tax_notes = data["tax_notes"]
    policies = data["policies"]
    doc = Document()
    configure(doc)
    add_cover(doc)

    add_para(doc, "目录", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True)
    for item, page in (
        ("一、学习引导", "3"), ("二、行业背景", "3"), ("三、企业资料", "3"),
        ("四、案头分析", "5"), ("五、约谈与证据核验", "8"), ("六、现场核验设计", "9"),
        ("七、评估结论", "9"), ("八、案例总结", "10"), ("九、案例大纲", "11"),
        ("附录A—D", "13"),
    ):
        add_para(doc, f"{item}……{page}", size=14)
    doc.add_page_break()

    add_heading(doc, 1, "一、学习引导")
    add_para(doc, "本案例面向软件和信息技术服务企业的税收风险管控场景，演示在仅取得两期公开年报、官方政策网页和行业统计资料时，如何建立“资料真实性—财务勾稽—风险线索—证据核验—结论分级”的可复现流程。研究不以会计波动代替税务结论；没有合同、发票、申报表、台账、银行流水、访谈或现场资料的事项，均保留为 TODO_MISSING_DATA。读者应掌握：公开财务数据可用于提出核验问题，但不能单独确认少缴税、骗取优惠或其他违法事实。")

    add_heading(doc, 1, "二、行业背景")
    add_para(doc, "软件和信息技术服务业常见业务涵盖软件产品、技术服务、实施交付及持续运维等形态。工信部发布的《2024年软件和信息技术服务业主要指标》显示，纳入重点企业统计的软件业务收入合计为137,276亿元、同比增长10.0%，其中信息技术服务收入92,190亿元、同比增长11.0%。行业收入确认、合同履约、预收款、研发归集和税收优惠资料之间具有较强关联，适合用跨表纵向观察定位待核验事项。上述行业数字仅作行业背景，不作为甲公司税务申报数据。")
    add_para(doc, "资料来源：[I1] 工业和信息化部《2024年软件和信息技术服务业主要指标》，2025-01-26；统计范围为重点企业，适用边界见该文注释。", size=8.5, color="666666")

    add_heading(doc, 1, "三、企业资料")
    add_heading(doc, 2, "（一）企业基本信息与匿名化边界")
    add_para(doc, "研究对象以“甲软件和信息技术服务企业”匿名呈现。提交文本不披露企业名称、统一社会信用代码、人员、学校或团队信息。S1、S2 为两期公开年度报告，其本地 PDF 与公开下载文件的 SHA-256 已逐份一致性核验；报告页数分别为279页、261页。企业成立时间、注册资本、经济性质、注册地址、增值税纳税人类型及纳税信用信息未在本研究的可用资料范围内，统一标注为 TODO_MISSING_DATA。")
    add_heading(doc, 2, "（二）企业经营情况")
    add_para(doc, "公开资料可将甲企业归入软件和信息技术服务行业，但未取得可逐项复核的产品清单、合同条款、客户结构、业务流程或开票流程资料。因此，本报告不对具体经营模式作事实性补述；上述业务资料均为 TODO_MISSING_DATA。")
    add_heading(doc, 2, "（三）两期会计报表摘录")
    base_rows = []
    labels = [
        ("operating_revenue", "营业收入"), ("operating_cost", "营业成本"), ("research_and_development_expense", "研发费用"),
        ("accounts_receivable", "应收账款"), ("contract_assets", "合同资产"), ("contract_liabilities", "合同负债"),
        ("government_grants_recognized_in_profit", "计入损益的政府补助"), ("taxes_paid", "支付的各项税费"),
    ]
    for field, label in labels:
        base_rows.append([
            label,
            format_million(amounts[2023][field]),
            format_million(amounts[2024][field]),
            f"{locations[2023][field]}；{locations[2024][field]}",
        ])
    add_table(doc, ["公开字段（单位：百万元）", "2023", "2024", "披露位置"], base_rows, [4.0, 2.0, 2.0, 6.7])
    add_para(doc, "注：金额均由公开年报文本提取程序生成；不对会计字段作纳税申报口径替代。", size=8.5, color="666666")
    add_para(doc, "注：本表为会计报表摘录而非完整资产负债表、利润表；完整报表附件由原始公开报告保存，不在匿名正文复印。", size=8.5, color="666666")
    add_heading(doc, 2, "（四）公开披露的税种及税费附注")
    tax_rows = []
    for label, key in (("应交增值税", "vat_payable"), ("应交企业所得税", "cit_payable"), ("应交个人所得税", "iit_payable"), ("税金及附加中的印花税", "stamp_tax_expense")):
        change = tax_notes[2024][key] / tax_notes[2023][key] - 1
        tax_rows.append([label, format_million(tax_notes[2023][key]), format_million(tax_notes[2024][key]), format_pct(change)])
    add_table(doc, ["项目（单位：百万元）", "2023", "2024", "同比变动"], tax_rows, [5.5, 2.1, 2.1, 3.0])
    add_para(doc, "应交税费为资产负债表日余额，税金及附加为会计费用列报。两者均不能直接等同于当期纳税申报额、应纳税额或应补税额。", size=8.5, color="666666")
    add_heading(doc, 2, "（五）纳税申报表")
    add_para(doc, "增值税、企业所得税及印花税的两期纳税申报表、附表和缴款书均未取得，状态为 TODO_MISSING_DATA。本报告不以公开年报、会计科目余额或模拟数据替代纳税申报表。")
    add_heading(doc, 2, "（六）案头资料")
    add_para(doc, "未取得经证实的行业正常区间、预警区间或税务风险等级口径，状态为 TODO_MISSING_DATA。为避免使用无法溯源的主观阈值，本案只呈现可复核的两期纵向观察，不对观察指标作高、中、低风险评级。")

    add_heading(doc, 1, "四、案头分析")
    add_heading(doc, 2, "（一）分析程序与结论规则")
    add_bullet(doc, "来源核验：公开年报 PDF 与公开下载文件进行 SHA-256 一致性核验；字段保留期间、报表页码、附注位置和代码血缘。")
    add_bullet(doc, "指标计算：采用“本期公开字段÷本期营业收入”的透明纵向比率；未设置经验阈值、行业阈值或主观风险分数。")
    add_bullet(doc, "结论门槛：OBSERVATION（描述性观察）→ ANOMALY（需有可复核异常规则）→ POTENTIAL_RISK → EVIDENCE_CHECK → FALSE_POSITIVE_CHECK → CONFIRMED_RISK。本案未越过 OBSERVATION。")
    add_bullet(doc, "反向核验：所有待核事项必须同时列明可能的正常业务解释和所需反证资料；缺失资料不以推定替代。")
    add_heading(doc, 2, "（二）十项公开资料观察指标")
    indicator_rows = []
    indicator_order = [f"PUB-IND-{number:02d}" for number in range(1, 11)]
    for indicator_id in indicator_order:
        current = indicators[indicator_id][2024]
        previous = indicators[indicator_id][2023]
        policy_state = current["policy_state"].replace("POLICY_", "")
        indicator_rows.append([
            indicator_id.replace("PUB-", ""), current["indicator_name"], format_pct(float(previous["value"])),
            format_pct(float(current["value"])), format_pct(float(current["change"])), policy_state,
        ])
    add_table(doc, ["编号", "指标", "2023", "2024", "变化", "政策状态"], indicator_rows, [1.4, 3.8, 1.7, 1.7, 1.7, 2.0])
    add_para(doc, "表中“政策状态”为政策层面的可核验状态，不是企业资格或实际税务处理状态。没有列出政策编号的综合财务指标统一标注为 POLICY_UNVERIFIED。", size=8.5, color="666666")
    add_heading(doc, 2, "（三）增值税：收入、回款与合同履约线索")
    add_para(doc, f"2024年销售回款/营业收入为{format_pct(float(indicators['PUB-IND-01'][2024]['value']))}，较2023年上升{format_pct(float(indicators['PUB-IND-01'][2024]['change']))}；应收账款/营业收入、合同资产/营业收入和合同负债/营业收入分别为{format_pct(float(indicators['PUB-IND-02'][2024]['value']))}、{format_pct(float(indicators['PUB-IND-03'][2024]['value']))}和{format_pct(float(indicators['PUB-IND-04'][2024]['value']))}。这些数字描述了公开披露的收款、应收和履约相关余额，不能确认增值税纳税义务发生时间、销项开票时间或申报结果。")
    add_para(doc, "政策核验：财税〔2016〕36号为营改增一般规则候选；财税〔2011〕100号涉及自行开发软件产品即征即退及进项税额分摊。政策原文已归档，但甲公司具体业务性质、软件产品资格、销售额拆分、进项分摊备案和申报资料均为 TODO_MISSING_DATA；适用性为 NEED_TAX_REVIEW。")
    add_table(doc, ["观察", "可能的正常解释", "必须取得的证据", "当前状态"], [[
        "收入、回款、应收、合同资产/负债的两期变动", "结算周期、项目验收、预收安排、履约进度、客户信用政策或合并范围变化", "合同、验收/履约资料、销项发票、增值税申报表、回款明细、期后回款、软件产品及进项分摊资料", "OBSERVATION；TODO_MISSING_DATA",
    ]], [3.4, 4.0, 5.4, 2.5])
    add_heading(doc, 2, "（四）企业所得税：研发费用与政府补助线索")
    add_para(doc, f"2024年研发费用/营业收入为{format_pct(float(indicators['PUB-IND-05'][2024]['value']))}，较2023年上升{format_pct(float(indicators['PUB-IND-05'][2024]['change']))}；计入损益政府补助/营业收入为{format_pct(float(indicators['PUB-IND-06'][2024]['value']))}，较2023年下降{format_pct(abs(float(indicators['PUB-IND-06'][2024]['change'])))}。这两项均为会计披露，不等于研发费用加计扣除金额或不征税收入金额。")
    add_para(doc, "政策核验：财政部、税务总局公告2023年第7号和第11号已核验其2023年度期间适用性；财税〔2011〕70号对专项用途财政性资金处理规定了条件。企业是否享受、项目是否符合、是否单独核算及申报处理均无公开足够证据。因此，不计算加计扣除额、应纳税所得额或补税额。")
    add_table(doc, ["观察", "可能的正常解释", "必须取得的证据", "当前状态"], [
        ["研发费用率上升", "研发投入结构和会计分类变化", "研发项目书、辅助账、人员工时、费用明细、A107012及企业所得税申报表", "OBSERVATION；TODO_MISSING_DATA"],
        ["政府补助金额及占比变动", "补助项目、确认时点、递延收益结转或财政性资金条件差异", "拨付文件、专项管理要求、单独核算资料、递延收益明细及企业所得税申报表", "OBSERVATION；TODO_MISSING_DATA"],
    ], [3.4, 4.0, 5.4, 2.5])
    add_heading(doc, 2, "（五）印花税：会计费用披露与凭证覆盖")
    stamp_change = tax_notes[2024]["stamp_tax_expense"] / tax_notes[2023]["stamp_tax_expense"] - 1
    add_para(doc, f"公开附注列示税金及附加中的印花税由2023年的{format_million(tax_notes[2023]['stamp_tax_expense'])}百万元变为2024年的{format_million(tax_notes[2024]['stamp_tax_expense'])}百万元，变动{format_pct(stamp_change)}。该变动只构成合同与税源台账核验的入口；公开资料没有逐份应税合同、应税凭证类别、计税金额、税源明细表、申报表或缴款书，故不能判断少缴、多缴、漏报或错用税目。")
    add_para(doc, "政策核验：国家税务总局公告2022年第14号已归档。应取得应税合同台账、印花税税源明细表、申报表和缴款书，并按凭证逐笔排除合同未生效、金额调整、免税或已完税等可能的假阳性因素。当前为 OBSERVATION；TODO_MISSING_DATA。")
    add_heading(doc, 2, "（六）综合税费观察")
    add_para(doc, f"2024年税金及附加/营业收入为{format_pct(float(indicators['PUB-IND-07'][2024]['value']))}，支付各项税费/营业收入为{format_pct(float(indicators['PUB-IND-08'][2024]['value']))}，应交税费/营业收入为{format_pct(float(indicators['PUB-IND-09'][2024]['value']))}。该等指标覆盖多税种、跨期结算和合并范围，未设单一税种阈值，政策状态为 POLICY_UNVERIFIED；仅保留为管理层案头勾稽观察。")

    add_heading(doc, 1, "五、约谈与证据核验")
    add_para(doc, "本研究未实施人员约谈，也未取得访谈纪要。为避免将未发生的程序写成既成事实，本节列示后续应开展的证据核验方案，状态为 PLANNED_NOT_PERFORMED，不构成访谈证据。")
    add_table(doc, ["税种/事项", "拟约谈角色", "核验问题", "留存证据"], [
        ["增值税", "财税负责人、业务负责人", "收入确认、开票、回款、软件产品及进项分摊的时间与口径是否一致？", "合同、验收单、发票清单、申报表、进项分摊备案/底稿"],
        ["企业所得税", "研发负责人、财税负责人", "研发项目、人员工时、辅助账、加计扣除和政府补助税务处理是否逐项勾稽？", "项目资料、辅助账、A107012、资金文件、单独核算资料"],
        ["印花税", "法务/合同管理员、财税负责人", "应税合同台账是否完整、合同变更和已完税凭证是否可逐笔对应？", "合同台账、税源明细表、申报表、缴款书、合同文本"],
    ], [2.2, 2.6, 5.4, 5.1])
    add_heading(doc, 2, "（一）假阳性排除清单")
    add_bullet(doc, "增值税线索：先按合同履约、验收、预收、开票、申报和回款时间线逐笔核验；正常结算周期或预收安排可解释公开余额变动。")
    add_bullet(doc, "研发线索：会计研发费用与税法可加计扣除费用存在范围、归集和留存资料差异；不得用会计金额直接乘优惠比例。")
    add_bullet(doc, "政府补助线索：会计确认与不征税收入条件不同；需同时核验财政性资金、专项用途、资金拨付文件和单独核算。")
    add_bullet(doc, "印花税线索：必须按应税凭证逐笔核验合同状态、计税依据、适用税目和完税凭证；费用变动本身不是风险证据。")

    add_heading(doc, 1, "六、现场核验设计")
    add_para(doc, "本研究未进入企业现场，未获取原始凭证或系统权限，也未形成现场工作底稿；状态为 NOT_PERFORMED。若后续取得授权，应以“抽样清单—原始凭证—系统记录—申报表—缴款书”闭环实施，不得以口头说明替代证据。")
    add_table(doc, ["核验模块", "现场动作", "通过标准", "未通过后的结论限制"], [
        ["收入与增值税", "抽取合同—验收—开票—申报—回款链路", "金额、期间、税率/方法和申报可勾稽", "最多推进为 POTENTIAL_RISK，待复核"],
        ["研发与企业所得税", "项目、人员、费用、辅助账、A107012逐项勾稽", "会计归集与税务申报差异有底稿解释", "不得计算优惠差额或补税额"],
        ["印花税", "合同台账与税源明细、申报、完税凭证抽样匹配", "应税凭证范围、计税依据及完税状态完整", "不得推定漏缴或确认风险"],
    ], [2.5, 5.0, 4.1, 3.7])

    add_heading(doc, 1, "七、评估结论")
    add_para(doc, "结论一：S1、S2 的公开年报文件已完成来源和哈希一致性核验；所列金额、比率和变动均由程序从同一来源数据生成，可复现。")
    add_para(doc, "结论二：本案形成了收入/回款/合同履约、研发费用、政府补助、综合税费及印花税费用等十项公开资料观察指标，覆盖增值税、企业所得税、印花税及其相关管理资料要求，期间为2023—2024年度。")
    add_para(doc, "结论三：由于没有取得纳税申报表、税款缴款书、合同、发票、台账、研发辅助账、政府补助拨付及单独核算资料，也未实施约谈和现场核验，全部事项止于 OBSERVATION 或 TODO_MISSING_DATA；本报告不存在 ANOMALY、POTENTIAL_RISK、EVIDENCE_CHECK、FALSE_POSITIVE_CHECK 或 CONFIRMED_RISK 的已完成结论。")
    add_para(doc, "结论四：本案不计算应补税额、滞纳金、罚款、责任人或风险等级。任何后续税务处理必须由取得完整证据后的专业人员，按纳税主体、业务事实、期间和当时有效政策重新核验。", bold=True)
    add_heading(doc, 2, "（一）拟定管控措施（未实施）")
    add_table(doc, ["管控对象", "建议控制点", "建议责任岗位", "建议时限", "复核证据"], [
        ["合同与增值税", "建立合同—履约—开票—申报—回款的期间勾稽清单", "业务、法务、财税岗位", "取得资料后10个工作日内完成首轮核验", "勾稽底稿、申报表、发票与合同抽样清单"],
        ["研发与企业所得税", "建立研发项目、人员工时、费用明细、辅助账与申报附表的对应关系", "研发、财税岗位", "取得资料后10个工作日内完成首轮核验", "项目资料、辅助账、A107012及差异说明"],
        ["印花税", "建立应税合同台账、税源明细、申报与缴款凭证的逐笔对应关系", "法务、合同管理、财税岗位", "取得资料后10个工作日内完成首轮核验", "合同台账、税源明细、申报表与缴款书"],
    ], [2.5, 4.6, 2.8, 3.1, 4.3])
    add_para(doc, "上述为未实施的建议方案，不是对甲企业现行岗位、内控或整改效果的事实陈述。", size=8.5, color="666666")

    add_heading(doc, 1, "八、案例总结")
    add_heading(doc, 2, "（一）案例简要描述")
    add_para(doc, "本案例以两期公开年报、官方政策原文和公开行业统计为输入，对软件和信息技术服务企业的收入履约、研发费用、政府补助、税费及印花税费用进行可复核的纵向观察，并逐项列明后续所需证据。")
    add_heading(doc, 2, "（二）案例分析特点")
    add_para(doc, "本案例的特点是建立可审计的公开信息研究链：原始文件哈希核验→字段提取→公式计算→政策来源归档→证据缺口清单→结论门槛控制。它避免了以会计数据替代税务申报、以政策标题替代适用性判断、以“没有数据”强行补造事实的错误。")
    add_heading(doc, 2, "（三）分析得失与适用边界")
    add_para(doc, "优点是金额、比率和数据来源可由同一程序复现；局限是无法取得原始业务和纳税资料，故不能确认税务风险、测算税额或评价整改效果。该研究适用于公开资料筛查和资料索取准备，不替代税务检查、审计或专业税务意见。")
    add_heading(doc, 2, "（四）后续监控建议")
    add_para(doc, "如获得企业授权，建议将本报告的十项观察逐项纳入周期性监控：每期更新公开字段；对变化项生成资料索取清单；只有完成证据核验和假阳性排除后，才可提出潜在风险或整改建议。没有授权或原始资料时，应保留本报告的公开资料边界。")

    add_heading(doc, 1, "九、案例大纲")
    add_table(doc, ["序号", "风险/指标名称", "税种", "期间", "当前阶段", "关键证据缺口"], [
        [str(index + 1), indicators[indicator_id][2024]["indicator_name"], indicators[indicator_id][2024]["candidate_tax_type"], "2023—2024", "OBSERVATION", indicators[indicator_id][2024]["evidence_required"]]
        for index, indicator_id in enumerate(indicator_order)
    ], [1.0, 3.4, 2.2, 1.8, 2.0, 7.0])

    doc.add_page_break()
    add_heading(doc, 1, "附录A　匿名化来源与政策核验索引")
    add_table(doc, ["编号", "资料", "核验信息/适用边界"], [
        ["S1", "甲公司2023年度公开报告", "279页；PDF哈希一致性已核验；仅支持公开财务披露，不替代税务申报或账簿。"],
        ["S2", "甲公司2024年度公开报告", "261页；PDF哈希一致性已核验；仅支持公开财务披露，不替代税务申报或账簿。"],
        ["I1", "工信部2024年软件和信息技术服务业主要指标", "发布于2025-01-26；用于行业背景，统计范围为重点企业。"],
    ], [1.2, 5.2, 9.0])
    policy_refs = [
        ("P1", "POL-VAT-2016-36"),
        ("P2", "POL-VAT-SOFTWARE-2011-100"),
        ("P3", "POL-CIT-RD-2023-07"),
        ("P4", "POL-CIT-RD-2023-11"),
        ("P5", "POL-CIT-GOV-2011-70"),
        ("P6", "POL-STAMP-2022-14"),
    ]
    policy_rows = []
    for ref, policy_id in policy_refs:
        policy = policies[policy_id]
        policy_rows.append([
            ref,
            f"{policy['document_number']}\n{policy['title']}",
            f"成文：{policy['publish_date']}\n生效：{policy['effective_date']}\n状态：{policy['status']}",
            f"条款：{policy['article']}\n范围：{policy['applicable_business']}\n边界：{policy['notes']}",
            policy["source_url"],
        ])
    add_heading(doc, 2, "（一）政策核验登记")
    add_table(doc, ["编号", "政策", "日期/状态", "条款、范围与核验边界", "官方来源"], policy_rows, [0.9, 3.1, 2.7, 6.0, 2.8])
    add_para(doc, "政策均采集自财政部、国家税务总局或其省级税务机关官方网站，核验日期为2026-09-05。提交版为匿名版，具体来源映射及原始文件哈希保存在研究底稿，不在正文披露企业名称。", size=8.5, color="666666")

    add_heading(doc, 1, "附录B　公式与数据血缘")
    formulas = []
    for indicator_id in indicator_order:
        row = indicators[indicator_id][2024]
        formulas.append([indicator_id.replace("PUB-", ""), row["indicator_name"], row["formula"], "2023、2024各期字段 / 各期营业收入", "无阈值；仅纵向描述"])
    add_table(doc, ["编号", "指标", "程序公式", "数据期间", "判定规则"], formulas, [1.4, 3.8, 4.4, 3.0, 3.0])
    add_para(doc, "生成路径：verify_public_disclosures.py → extract_public_financial_fields.py → screen_public_indicators.py → build_submission_document.py。提交文本中的核心金额、比例和变动均由此路径输出。", size=8.5, color="666666")

    add_heading(doc, 1, "附录C　资料索取与结论升级门槛")
    add_table(doc, ["阶段", "最小条件", "本案状态"], [
        ["OBSERVATION", "可复核公开字段及公式", "已完成"],
        ["ANOMALY", "预先定义的异常规则和可复核结果", "未进入；未设主观阈值"],
        ["POTENTIAL_RISK", "具体业务事实与适用政策的初步对应", "未进入；TODO_MISSING_DATA"],
        ["EVIDENCE_CHECK", "合同、发票、申报表、账簿等原始证据", "未完成"],
        ["FALSE_POSITIVE_CHECK", "对正常业务解释的逐项反证", "未完成"],
        ["CONFIRMED_RISK", "完整事实、税法适用、金额计算及复核", "不存在"],
    ], [3.2, 9.3, 5.0])

    add_heading(doc, 1, "附录D　提交合规自检")
    add_table(doc, ["检查项", "结果", "说明"], [
        ["企业、人员、学校和团队匿名", "通过", "仅使用甲公司、S1/S2等匿名标识。"],
        ["两期数据与三类税种/事项", "通过", "覆盖2023—2024年；增值税、企业所得税、印花税。"],
        ["十项指标", "通过", "十项均为公开资料观察，未虚构风险点。"],
        ["税法来源及适用边界", "通过", "附录A列明官方政策与需复核事项。"],
        ["访谈和现场材料", "未提供", "如实标注 NOT_PERFORMED，不以虚构程序替代。"],
        ["确认风险及税额", "不存在", "无原始证据，不作确认或金额推定。"],
    ], [4.5, 2.5, 10.5])
    return doc


tracker = install_script_logging(
    "build_submission_document",
    "python scripts/build_submission_document.py",
    "PHASE 1",
    ROOT,
    [DOCX.relative_to(ROOT), PDF.relative_to(ROOT)],
)
for required in (FIELDS, INDICATORS, POLICIES, *REPORT_TEXTS.values()):
    if not required.exists():
        raise FileNotFoundError(f"required input missing: {required}")
OUT.mkdir(parents=True, exist_ok=True)
document = build_document(collect_data())
document.save(DOCX)
tracker.note("DOCX rendered from verified public-disclosure outputs")
subprocess.run(
    ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUT), str(DOCX)],
    check=True,
    cwd=ROOT,
    env={**os.environ, "FONTCONFIG_FILE": str(FONT_CONFIG)},
)
if not PDF.exists() or PDF.stat().st_size == 0:
    raise RuntimeError("PDF conversion did not produce a file")
tracker.note("PDF rendered from DOCX by LibreOffice")
print(f"created {DOCX}")
print(f"created {PDF}")
